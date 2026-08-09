"""Génère un CV au format DOCX (facilement modifiable / exportable en PDF
depuis Word) à partir du même contexte que le template HTML."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from cv_generator import DEFAULT_CV_LABELS


BLUE = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x4B, 0x55, 0x63)


def _fmt(d, fmt):
    return d.strftime(fmt) if d else ""


def build_docx(context: dict) -> Document:
    profile = context["profile"]
    labels = context.get("labels", DEFAULT_CV_LABELS)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    run = title.add_run(profile.full_name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    contact_bits = [b for b in [profile.email, profile.phone, profile.location,
                                 profile.linkedin, profile.github, profile.website] if b]
    if contact_bits:
        p = doc.add_paragraph(" | ".join(contact_bits))
        p.runs[0].font.color.rgb = GREY
        p.runs[0].font.size = Pt(9.5)

    if context.get("summary"):
        p = doc.add_paragraph(context["summary"])
        p.runs[0].italic = True

    def add_heading(text):
        h = doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.color.rgb = BLUE
            run.font.size = Pt(12)

    if context.get("experiences"):
        add_heading(labels["experiences"])
        for e in context["experiences"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{e['role']} — {e['company']}")
            r.bold = True
            dates = f"{_fmt(e['start_date'], '%m/%Y')} – {_fmt(e['end_date'], '%m/%Y') or labels['current']}"
            p.add_run(f"    ({dates})").italic = True
            if e.get("location"):
                loc = doc.add_paragraph(e["location"])
                loc.runs[0].font.color.rgb = GREY
                loc.runs[0].font.size = Pt(9.5)
            for b in e.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    if context.get("skills_by_category"):
        add_heading(labels["skills"])
        for cat, items in context["skills_by_category"].items():
            p = doc.add_paragraph()
            p.add_run(f"{cat} : ").bold = True
            names = ", ".join(
                (s["name"] + (f" ({s['level']})" if s.get("level") else "")) for s in items
            )
            p.add_run(names)

    if context.get("projects"):
        add_heading(labels["projects"])
        for pr in context["projects"]:
            p = doc.add_paragraph()
            p.add_run(pr["name"]).bold = True
            if pr.get("link"):
                p.add_run(f"  ({pr['link']})").italic = True
            if pr.get("description"):
                doc.add_paragraph(pr["description"])

    if context.get("educations"):
        add_heading(labels["education"])
        for ed in context["educations"]:
            p = doc.add_paragraph()
            label = ed.degree + (f" — {ed.field}" if ed.field else "") + f", {ed.school}"
            r = p.add_run(label)
            r.bold = True
            years = f"{_fmt(ed.start_date, '%Y')} – {_fmt(ed.end_date, '%Y') or labels['current']}"
            p.add_run(f"    ({years})").italic = True
            if ed.description:
                doc.add_paragraph(ed.description)

    if context.get("languages"):
        add_heading(labels["languages"])
        doc.add_paragraph(" | ".join(f"{l.name} — {l.level}" for l in context["languages"]))

    return doc


def build_letter_docx(context: dict) -> Document:
    profile = context["profile"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run(profile.full_name)
    r.bold = True
    r.font.size = Pt(13)

    contact_bits = [b for b in [profile.email, profile.phone, profile.location] if b]
    if contact_bits:
        cp = doc.add_paragraph(" · ".join(contact_bits))
        cp.runs[0].font.color.rgb = GREY
        cp.runs[0].font.size = Pt(9.5)

    date_p = doc.add_paragraph(f"Le {context['today']}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    obj = doc.add_paragraph()
    obj_text = "Objet : Candidature"
    if context.get("job_title"):
        obj_text += f" au poste de {context['job_title']}"
    obj.add_run(obj_text).bold = True

    doc.add_paragraph("Madame, Monsieur,")

    intro = context.get("summary") or "Je vous adresse ma candidature avec grand intérêt."
    if context.get("job_title"):
        intro += (
            f" C'est donc naturellement que je souhaite vous soumettre ma candidature "
            f"pour le poste de {context['job_title']}."
        )
    doc.add_paragraph(intro)

    if context.get("top_experiences"):
        exps = ", ainsi que ".join(
            f"en tant que {e.role} chez {e.company}" for e in context["top_experiences"]
        )
        doc.add_paragraph(
            f"Mon parcours m'a permis de développer une expérience concrète, notamment {exps}, "
            f"où j'ai pu mettre en pratique des compétences directement utiles au poste que "
            f"vous proposez."
        )

    if context.get("top_skills"):
        doc.add_paragraph(
            f"Je maîtrise notamment {', '.join(context['top_skills'])}, des compétences que je "
            f"serais heureux(se) de mettre au service de votre équipe."
        )

    doc.add_paragraph(
        "Je me tiens à votre disposition pour un entretien au cours duquel je pourrais vous "
        "exposer plus en détail ma motivation et mon parcours. Je vous remercie par avance "
        "pour l'attention portée à ma candidature."
    )
    doc.add_paragraph(
        "Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées."
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(profile.full_name).bold = True

    return doc


def build_letter_docx_ai(context: dict) -> Document:
    """Lettre de motivation en mode IA : contrairement à build_letter_docx, aucune
    phrase française n'est codée en dur ici -- tout le texte (objet, formule
    d'appel, corps, formule de politesse) vient déjà rédigé/traduit par l'IA."""
    profile = context["profile"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run(profile.full_name)
    r.bold = True
    r.font.size = Pt(13)

    contact_bits = [b for b in [profile.email, profile.phone, profile.location] if b]
    if contact_bits:
        cp = doc.add_paragraph(" · ".join(contact_bits))
        cp.runs[0].font.color.rgb = GREY
        cp.runs[0].font.size = Pt(9.5)

    date_p = doc.add_paragraph(context["today"])
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    obj = doc.add_paragraph()
    obj.add_run(context["subject_line"]).bold = True

    doc.add_paragraph(context["salutation"])

    for paragraph in context["paragraphs"]:
        doc.add_paragraph(paragraph)

    doc.add_paragraph(context["closing"])

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(profile.full_name).bold = True

    return doc
